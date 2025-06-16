"""
Utility per l'estrazione del testo da diversi formati di documenti.
"""
import os
import logging
import magic
from typing import Optional, List, Tuple
from pathlib import Path
from django.conf import settings
import PyPDF2
import docx
import pandas as pd
from PIL import Image
import pytesseract

logger = logging.getLogger(__name__)

class TextExtractor:
    """
    Classe per estrarre testo da diversi formati di documenti.
    """
    
    SUPPORTED_FORMATS = {
        'application/pdf': 'pdf',
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx',
        'application/msword': 'doc',
        'text/plain': 'txt',
        'image/jpeg': 'image',
        'image/png': 'image',
        'image/tiff': 'image',
        'image/bmp': 'image',
        'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': 'xlsx',
        'application/vnd.ms-excel': 'xls',
    }
    
    def __init__(self):
        """
        Inizializza l'estrattore di testo.
        """
        self.chunk_size = 1000
        self.chunk_overlap = 200
    
    def detect_file_type(self, file_path: str) -> str:
        """
        Rileva il tipo di file usando python-magic.
        
        Args:
            file_path (str): Percorso del file
            
        Returns:
            str: MIME type del file
        """
        try:
            mime_type = magic.from_file(file_path, mime=True)
            logger.info(f"Rilevato tipo di file: {mime_type} per {file_path}")
            return mime_type
        except Exception as e:
            logger.error(f"Errore nel rilevamento del tipo di file {file_path}: {str(e)}")
            # Fallback basato sull'estensione
            ext = Path(file_path).suffix.lower()
            mime_mappings = {
                '.pdf': 'application/pdf',
                '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingtml.document',
                '.doc': 'application/msword',
                '.txt': 'text/plain',
                '.jpg': 'image/jpeg',
                '.jpeg': 'image/jpeg',
                '.png': 'image/png',
                '.tiff': 'image/tiff',
                '.bmp': 'image/bmp',
                '.xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                '.xls': 'application/vnd.ms-excel',
            }
            return mime_mappings.get(ext, 'application/octet-stream')
    
    def is_supported_format(self, file_path: str) -> bool:
        """
        Verifica se il formato del file è supportato.
        
        Args:
            file_path (str): Percorso del file
            
        Returns:
            bool: True se il formato è supportato
        """
        mime_type = self.detect_file_type(file_path)
        return mime_type in self.SUPPORTED_FORMATS
    
    def extract_text_from_pdf(self, file_path: str) -> str:
        """
        Estrae testo da file PDF.
        
        Args:
            file_path (str): Percorso del file PDF
            
        Returns:
            str: Testo estratto
        """
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text += page.extract_text()
                
                logger.info(f"Estratto testo da PDF: {len(text)} caratteri")
                return text.strip()
                
        except Exception as e:
            logger.error(f"Errore nell'estrazione da PDF {file_path}: {str(e)}")
            raise Exception(f"Errore nell'estrazione da PDF: {str(e)}")
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """
        Estrae testo da file DOCX.
        
        Args:
            file_path (str): Percorso del file DOCX
            
        Returns:
            str: Testo estratto
        """
        try:
            doc = docx.Document(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Estrai anche il testo dalle tabelle
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            logger.info(f"Estratto testo da DOCX: {len(text)} caratteri")
            return text.strip()
            
        except Exception as e:
            logger.error(f"Errore nell'estrazione da DOCX {file_path}: {str(e)}")
            raise Exception(f"Errore nell'estrazione da DOCX: {str(e)}")
    
    def extract_text_from_txt(self, file_path: str) -> str:
        """
        Estrae testo da file TXT.
        
        Args:
            file_path (str): Percorso del file TXT
            
        Returns:
            str: Testo estratto
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            
            logger.info(f"Estratto testo da TXT: {len(text)} caratteri")
            return text.strip()
            
        except UnicodeDecodeError:
            # Prova con encoding diversi
            encodings = ['latin-1', 'iso-8859-1', 'cp1252']
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        text = file.read()
                    logger.info(f"Estratto testo da TXT con encoding {encoding}: {len(text)} caratteri")
                    return text.strip()
                except UnicodeDecodeError:
                    continue
            raise Exception("Impossibile decodificare il file di testo")
        except Exception as e:
            logger.error(f"Errore nell'estrazione da TXT {file_path}: {str(e)}")
            raise Exception(f"Errore nell'estrazione da TXT: {str(e)}")
    
    def extract_text_from_image(self, file_path: str) -> str:
        """
        Estrae testo da immagine usando OCR.
        
        Args:
            file_path (str): Percorso del file immagine
            
        Returns:
            str: Testo estratto
        """
        try:
            image = Image.open(file_path)
            text = pytesseract.image_to_string(image, lang='ita+eng')
            
            logger.info(f"Estratto testo da immagine: {len(text)} caratteri")
            return text.strip()
            
        except Exception as e:
            logger.error(f"Errore nell'estrazione da immagine {file_path}: {str(e)}")
            raise Exception(f"Errore nell'estrazione da immagine: {str(e)}")
    
    def extract_text_from_excel(self, file_path: str) -> str:
        """
        Estrae testo da file Excel.
        
        Args:
            file_path (str): Percorso del file Excel
            
        Returns:
            str: Testo estratto
        """
        try:
            # Leggi tutti i fogli
            excel_data = pd.read_excel(file_path, sheet_name=None)
            text = ""
            
            for sheet_name, df in excel_data.items():
                text += f"--- {sheet_name} ---\n"
                # Converti DataFrame in testo
                text += df.to_string(index=False, na_rep='')
                text += "\n\n"
            
            logger.info(f"Estratto testo da Excel: {len(text)} caratteri")
            return text.strip()
            
        except Exception as e:
            logger.error(f"Errore nell'estrazione da Excel {file_path}: {str(e)}")
            raise Exception(f"Errore nell'estrazione da Excel: {str(e)}")
    
    def extract_text(self, file_path: str) -> str:
        """
        Estrae testo dal file basandosi sul suo tipo.
        
        Args:
            file_path (str): Percorso del file
            
        Returns:
            str: Testo estratto
            
        Raises:
            Exception: Se il formato non è supportato o si verifica un errore
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File non trovato: {file_path}")
        
        mime_type = self.detect_file_type(file_path)
        
        if mime_type not in self.SUPPORTED_FORMATS:
            raise Exception(f"Formato file non supportato: {mime_type}")
        
        format_type = self.SUPPORTED_FORMATS[mime_type]
        
        try:
            if format_type == 'pdf':
                return self.extract_text_from_pdf(file_path)
            elif format_type == 'docx':
                return self.extract_text_from_docx(file_path)
            elif format_type == 'txt':
                return self.extract_text_from_txt(file_path)
            elif format_type == 'image':
                return self.extract_text_from_image(file_path)
            elif format_type in ['xlsx', 'xls']:
                return self.extract_text_from_excel(file_path)
            else:
                raise Exception(f"Handler non implementato per il formato: {format_type}")
                
        except Exception as e:
            logger.error(f"Errore nell'estrazione del testo da {file_path}: {str(e)}")
            raise
    
    def chunk_text(self, text: str, chunk_size: int = None, overlap: int = None) -> List[str]:
        """
        Divide il testo in chunk di dimensione specifica con overlap.
        
        Args:
            text (str): Testo da dividere
            chunk_size (int): Dimensione del chunk (default: self.chunk_size)
            overlap (int): Sovrapposizione tra chunk (default: self.chunk_overlap)
            
        Returns:
            List[str]: Lista di chunk di testo
        """
        if chunk_size is None:
            chunk_size = self.chunk_size
        if overlap is None:
            overlap = self.chunk_overlap
        
        if len(text) <= chunk_size:
            return [text]
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + chunk_size
            
            if end < len(text):
                # Cerca il punto di divisione migliore (fine di frase o paragrafo)
                chunk = text[start:end]
                
                # Trova l'ultimo punto, new line o spazio
                last_sentence = max(
                    chunk.rfind('.'),
                    chunk.rfind('\n'),
                    chunk.rfind('!'),
                    chunk.rfind('?')
                )
                
                if last_sentence > chunk_size * 0.5:  # Solo se non troppo piccolo
                    chunk = text[start:start + last_sentence + 1]
                    end = start + last_sentence + 1
                else:
                    # Trova l'ultimo spazio
                    last_space = chunk.rfind(' ')
                    if last_space > chunk_size * 0.5:
                        chunk = text[start:start + last_space]
                        end = start + last_space
            else:
                chunk = text[start:]
            
            chunks.append(chunk.strip())
            
            if end >= len(text):
                break
                
            start = end - overlap
        
        logger.info(f"Testo diviso in {len(chunks)} chunk")
        return chunks

def extract_text(file_path: str) -> Optional[str]:
    """
    Estrae il testo da un file supportando vari formati.
    """
    try:
        file_extension = os.path.splitext(file_path)[1].lower()
        
        if file_extension == '.pdf':
            return extract_from_pdf(file_path)
        elif file_extension == '.docx':
            return extract_from_docx(file_path)
        elif file_extension == '.txt':
            return extract_from_txt(file_path)
        elif file_extension in ['.csv', '.xlsx', '.xls']:
            return extract_from_spreadsheet(file_path)
        else:
            logger.error(f"Formato file non supportato: {file_extension}")
            return None
            
    except Exception as e:
        logger.error(f"Errore nell'estrazione del testo da {file_path}: {str(e)}")
        return None

def extract_from_pdf(file_path: str) -> str:
    """
    Estrae il testo da un file PDF.
    """
    text = ""
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        return text
    except Exception as e:
        logger.error(f"Errore nell'estrazione del testo dal PDF {file_path}: {str(e)}")
        raise

def extract_from_docx(file_path: str) -> str:
    """
    Estrae il testo da un file DOCX.
    """
    try:
        doc = docx.Document(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    except Exception as e:
        logger.error(f"Errore nell'estrazione del testo dal DOCX {file_path}: {str(e)}")
        raise

def extract_from_txt(file_path: str) -> str:
    """
    Estrae il testo da un file di testo.
    """
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    except Exception as e:
        logger.error(f"Errore nell'estrazione del testo dal file di testo {file_path}: {str(e)}")
        raise

def extract_from_spreadsheet(file_path: str) -> str:
    """
    Estrae il testo da un file Excel o CSV.
    """
    try:
        if file_path.endswith('.csv'):
            df = pd.read_csv(file_path)
        else:
            df = pd.read_excel(file_path)
        
        # Converti il DataFrame in testo
        text = df.to_string(index=False)
        return text
    except Exception as e:
        logger.error(f"Errore nell'estrazione del testo dal foglio di calcolo {file_path}: {str(e)}")
        raise 